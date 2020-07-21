from docx import Document
import pandas as pd
import re
import math
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
def lookup_fn2(df, key_col, key_val):
    try:
        result_df = df.loc[df[key_col] == key_val].iloc[0]
        return result_df
    except IndexError:
        return pd.DataFrame()
#Get Dict for valid AIPs
def GetValidAIPForCollection(origDf, collectionCSVNamesToName):
    print("new")
    bagRegex = r"^Bag-(.*)_(.*).zip$"
    bagRegex2 = r"^(.*)_(.*)_foxml_atomzip.zip$"
    df = pd.read_csv("bags_validation_report.csv")
    numBagPID = df.shape[0]
    numObjects = origDf.shape[0]
    numValid = 0
    numInvalid = 0
    totalAIP = 0
    totalValidAIP = 0
    pidCol = origDf["PID"]
    collectionCol = origDf["isMemberOfCollection"]
    collectionToNumValid = {}
    for index, row in df.iterrows():
        bag = row["bag_name"]
        isBagValid = False
        if re.match(bagRegex, bag):
            pid = re.search(bagRegex, bag).group(1)+":"+re.search(bagRegex, bag).group(2)
            isBagValid = True
        elif re.match(bagRegex2, bag):
            pid = re.search(bagRegex2, bag).group(1)+":"+re.search(bagRegex2, bag).group(2)
            isBagValid = True
        if isBagValid:
            origRow = lookup_fn2(origDf, "PID", pid)
            if (not origRow.empty):
                collectionRaw = origRow["isMemberOfCollection"]
                if pd.notna(collectionRaw) and pd.isna(origRow["isPageOf"]) and pd.isna(origRow["isConstituentOf"]):
                    collection = collectionCSVNamesToName[collectionRaw]
                    if collection not in collectionToNumValid:
                        collectionToNumValid[collection] = [0,0]
                    collectionToNumValid[collection][0]+=1
                    totalAIP+=1
                    valid = row["validation_status"]
                    if valid:
                        collectionToNumValid[collection][1]+=1
                        totalValidAIP+=1
            else:
                print("not found",pid)
    return collectionToNumValid, totalAIP, totalValidAIP

def GenerateBarPlot(xValues, yValues, fileName, header):
    N = len(xValues)
    plt.figure(figsize=(20, 7))  # width:20, height:7
    plt.bar(range(N), yValues, align='edge', width=0.6)
    plt.xticks(range(N), xValues,rotation=90)

    plt.gcf().subplots_adjust(bottom=0.3)
    plt.xticks(range(N)) # add loads of ticks
    plt.grid()

    plt.gca().margins(x=0)
    plt.gcf().canvas.draw()
    tl = plt.gca().get_xticklabels()
    maxsize = max([t.get_window_extent().width for t in tl])
    m = 2.0 # inch margin
    s = maxsize/plt.gcf().dpi*N+2*m
    margin = m/plt.gcf().get_size_inches()[0]

    plt.gcf().subplots_adjust(left=margin, right=1.-margin)
    plt.gcf().set_size_inches(s, plt.gcf().get_size_inches()[1])
    plt.title(header)
    plt.savefig(fileName)

def GenerateLinePlot(xValues, yValues, fileName, header):
    N = len(xValues)
    plt.figure(figsize=(20, 7))  # width:20, height:7
    plt.plot(range(N), yValues,'--bo')
    plt.xticks(range(N), xValues,rotation=90)

    plt.gcf().subplots_adjust(bottom=0.3)
    plt.xticks(range(N)) # add loads of ticks
    plt.grid()

    plt.gca().margins(x=0)
    plt.gcf().canvas.draw()
    tl = plt.gca().get_xticklabels()
    maxsize = max([t.get_window_extent().width for t in tl])
    m = 2.0 # inch margin
    s = maxsize/plt.gcf().dpi*N+2*m
    margin = m/plt.gcf().get_size_inches()[0]

    plt.gcf().subplots_adjust(left=margin, right=1.-margin)
    plt.gcf().set_size_inches(s, plt.gcf().get_size_inches()[1])
    plt.title(header)
    plt.savefig(fileName)

def GenerateTable(leftColName, rightColName, data, table):
    table.style = 'Table Grid'
    i = 1
    row = table.rows[0]
    row.cells[0].paragraphs[0].add_run(leftColName).bold = True
    row.cells[1].paragraphs[0].add_run(rightColName).bold = True
    for key, value in sorted(data.items()):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = str(value)
        i+=1

def GenerateTitle(document, title, numDiagram, expText = ""):
    p = document.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = document.styles['Header']
    if expText != "":
        p = document.add_paragraph(expText)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(numDiagram)
